import re
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from loguru import logger
import asyncio
from concurrent.futures import ThreadPoolExecutor
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Word document processing will be disabled.")

try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False


class WordDocumentProcessor:
    
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=2)
        self.nlp = None
        self._load_spacy_model()
        
    def _load_spacy_model(self):
        if not SPACY_AVAILABLE:
            logger.warning("spaCy not available. Using regex-based processing.")
            self.nlp = None
            return
            
        try:
            try:
                self.nlp = spacy.load("en_core_web_trf")
                logger.info("Loaded spaCy transformer model")
            except OSError:
                self.nlp = spacy.load("en_core_web_sm")
                logger.info("Loaded spaCy small model")
        except OSError:
            logger.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
    
    async def process_word_document(
        self,
        doc_path: str,
        paper_type: str = "auto"
    ) -> Dict[str, Any]:
        try:
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library not available. Please install it with: pip install python-docx")
            
            references = await asyncio.get_event_loop().run_in_executor(
                self.executor,
                self._extract_references_from_docx,
                doc_path
            )
            
            paper_data = await asyncio.get_event_loop().run_in_executor(
                self.executor,
                self._extract_paper_metadata,
                doc_path
            )
            
            return {
                "success": True,
                "paper_data": paper_data,
                "references": references,
                "reference_count": len(references),
                "paper_type": paper_type
            }
                
        except Exception as e:
            logger.error(f"Word document processing error: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "references": [],
                "reference_count": 0
            }
    
    def _extract_references_from_docx(self, doc_path: str) -> List[Dict[str, Any]]:
        references = []
        
        try:
            doc = Document(doc_path)
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
                full_text += "\n"
            
            logger.info(f"Word document text extracted: {len(full_text)} characters")
            
            if full_text:
                references = self._extract_references_from_text(full_text)
                logger.info(f"Word document processing completed: {len(references)} references found")
            else:
                logger.warning("No text extracted from Word document")
                    
        except Exception as e:
            logger.error(f"Error extracting references from Word document: {e}")
        
        return references
    
    def _extract_references_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Extract references from text using simplified approach focusing on document end"""
        references = []
        
        # Split text into lines
        lines = text.split('\n')
        total_lines = len(lines)
        
        logger.info(f"Processing {total_lines} lines from Word document")
        
        # Strategy 1: Look for reference section in last 20% of text (or last 1-2 pages)
        search_start = max(0, int(total_lines * 0.8))
        logger.info(f"Searching for references starting from line {search_start} (last 20% of {total_lines} lines)")
        
        
        ref_section_start = self._find_reference_section(lines[search_start:])
        if ref_section_start is not None:
            actual_start = search_start + ref_section_start
            logger.info(f"Found reference section at line {actual_start}")
            references = self._extract_references_from_section(lines[actual_start:])
        
        if len(references) < 5:
            logger.info(f"  Only found {len(references)} references, expanding search to last 30%")
            search_start = max(0, int(total_lines * 0.7))
            ref_section_start = self._find_reference_section(lines[search_start:])
            if ref_section_start is not None:
                actual_start = search_start + ref_section_start
                logger.info(f"Found reference section at line {actual_start}")
                references = self._extract_references_from_section(lines[actual_start:])
        
        if len(references) < 5:
            logger.info(f"  Only found {len(references)} references, trying last 1-2 pages")
            last_pages_start = max(0, total_lines - 100)
            ref_section_start = self._find_reference_section(lines[last_pages_start:])
            if ref_section_start is not None:
                actual_start = last_pages_start + ref_section_start
                logger.info(f"Found reference section at line {actual_start}")
                references = self._extract_references_from_section(lines[actual_start:])
        
        if len(references) < 3:
            logger.info(f"  Only found {len(references)} references, trying numbered reference extraction on entire document")
            numbered_refs = self._extract_numbered_references(text)
            if numbered_refs:
                structured_refs = []
                for ref_text in numbered_refs:
                    structured_refs.append({
                        "raw": ref_text.strip()
                    })
                references = structured_refs
                logger.info(f" Found {len(references)} numbered references in entire document")
        
        if len(references) < 3:
            logger.info(f"  Only found {len(references)} references, searching entire document for reference section")
            ref_section_start = self._find_reference_section(lines)
            if ref_section_start is not None:
                logger.info(f"Found reference section at line {ref_section_start}")
                references = self._extract_references_from_section(lines[ref_section_start:])
        
        if len(references) < 3:
            logger.info(f"  Only found {len(references)} references, trying aggressive pattern matching")
            references = self._extract_references_aggressive(text)
        
        if len(references) < 3:
            logger.info(f"  Only found {len(references)} references, trying very loose pattern matching")
            references = self._extract_references_loose(text)
        
        logger.info(f" Reference extraction completed: {len(references)} references found")
        return references
    
    def _find_reference_section(self, lines: List[str]) -> Optional[int]:
        """Find the start of reference section in given lines"""
        ref_keywords = ['references', 'reference', 'bibliography', 'works cited', 'literature cited', 'citations']
        
        exclude_keywords = ['related works', 'related work', 'literature review', 'background', 'introduction', 'methodology', 'conclusion']
        
        logger.info(f" Searching for reference section in {len(lines)} lines")
        
        for i, line in enumerate(lines):
            line_clean = line.strip().lower()
            original_line = line.strip()
            
            if i < 10:
                logger.info(f" Line {i}: '{original_line}'")
            
            for exclude_keyword in exclude_keywords:
                if exclude_keyword in line_clean:
                    logger.info(f"Skipping '{original_line}' - contains exclude keyword: {exclude_keyword}")
                    continue
            
            for keyword in ref_keywords:
                if (keyword in line_clean and len(line_clean.split()) <= 4) or \
                   re.match(rf'^{keyword}\s*:?\s*$', line_clean):
                    logger.info(f" Found reference section header: '{original_line}'")
                    return i
            
            if re.match(r'^\d+[\.\)]\s*(references?|bibliography|citations?)', line_clean):
                logger.info(f" Found numbered reference section: '{original_line}'")
                return i
            
            if re.match(r'^(reference|references)\s*:?\s*$', line_clean):
                logger.info(f" Found standalone reference header: '{original_line}'")
                return i
        
        logger.info(f" No reference section found in {len(lines)} lines")
        return None
    
    def _extract_references_from_section(self, lines: List[str]) -> List[Dict[str, Any]]:
        """Extract individual references from reference section - SIMPLE SPLITTING ONLY"""
        full_text = "\n".join(lines)
        
        logger.info(f" Processing {len(lines)} lines from reference section")
        
        references = self._extract_references_simple(full_text)
        
        structured_refs = []
        for i, ref_text in enumerate(references):
            structured_refs.append({
                "raw": ref_text.strip()
            })
        
        logger.info(f" Extracted {len(structured_refs)} references")
        return structured_refs
    
    def _extract_references_simple(self, text: str) -> List[str]:
        """Simple, reliable reference extraction using regex patterns"""
        logger.info(" Using simple reference extraction")
        
        references = []
        
        # Split on bracket references first (most reliable) - handle newlines properly
        lines = text.split('\n')
        
        logger.info(f" Processing {len(lines)} lines for bracket references")
        
        current_ref = ""
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if re.match(r'^\[\d+\]', line):
                line_lower = line.lower()
                is_related_work = any(indicator in line_lower for indicator in [
                    'this paper', 'our approach', 'we present', 'we propose'
                ])
                
                if not is_related_work:
                    if current_ref:
                        references.append(current_ref.strip())
                    current_ref = line
            elif current_ref:
                current_ref += " " + line
            elif len(line) > 20:
                current_ref = line
        
        if current_ref:
            references.append(current_ref.strip())
        
        logger.info(f" Bracket extraction found {len(references)} references")
        
        if len(references) < 3:
            logger.info("  Few bracket references found, trying numbered references")
            references = self._extract_numbered_references(text)
        
        if len(references) < 3:
            logger.info("  Few numbered references found, trying author-year format")
            references = self._extract_author_year_references(text)
        
        logger.info(f" Simple extraction found {len(references)} references")
        return references
    
    def _extract_numbered_references(self, text: str) -> List[str]:
        """Extract numbered references (1. Author, 2) Author, etc.) - handles newlines properly"""
        references = []
        
        lines = text.split('\n')
        
        logger.info(f" Processing {len(lines)} lines for numbered references")
        
        exclude_indicators = [
            'related works', 'related work', 'literature review', 'background',
            'introduction', 'methodology', 'conclusion', 'discussion'
        ]
        
        current_ref = ""
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            line_lower = line.lower()
            
            is_related_work = False
            for indicator in exclude_indicators:
                if indicator in line_lower:
                    logger.info(f"Skipping line - contains related work indicator: {line[:100]}...")
                    is_related_work = True
                    break
            
            if is_related_work:
                continue
            
            if re.match(r'^\d+\.\s*', line) and not re.match(r'^\[\d+\]', line):
                if current_ref:
                    references.append(current_ref.strip())
                    logger.info(f" Found numbered reference: {current_ref[:100]}...")
                current_ref = line
            elif current_ref:
                current_ref += " " + line
            elif len(line) > 30:
                if not any(indicator in line_lower for indicator in exclude_indicators):
                    current_ref = line
        
        if current_ref:
            references.append(current_ref.strip())
            logger.info(f" Found final numbered reference: {current_ref[:100]}...")
        
        logger.info(f" Numbered extraction found {len(references)} references")
        return references
    
    def _extract_author_year_references(self, text: str) -> List[str]:
        """Extract author-year references (Smith, J. (2020), etc.) - handles newlines properly"""
        references = []
        
        lines = text.split('\n')
        
        current_ref = ""
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if re.match(r'^[A-Z][a-z]+,\s*[A-Z]\.\s*\(\d{4}\)', line):
                if current_ref:
                    references.append(current_ref.strip())
                current_ref = line
            elif current_ref:
                current_ref += " " + line
            elif len(line) > 20:
                current_ref = line
        
        if current_ref:
            references.append(current_ref.strip())
        
        return references
    
    def _extract_references_aggressive(self, text: str) -> List[Dict[str, Any]]:
        """Aggressive reference extraction that looks for common reference patterns anywhere in the document"""
        references = []
        
        logger.info(" Using aggressive reference extraction")
        
        lines = text.split('\n')
        
        # Look for any line that might be a reference based on common patterns
        potential_refs = []
        
        exclude_indicators = [
            'related works', 'related work', 'literature review', 'background',
            'introduction', 'methodology', 'conclusion', 'discussion'
        ]
        
        for line in lines:
            line = line.strip()
            if not line or len(line) < 20:  # Lowered threshold
                continue
            
            line_lower = line.lower()
            
            is_related_work = False
            for indicator in exclude_indicators:
                if indicator in line_lower:
                    logger.info(f"Skipping potential reference - contains related work indicator: {line[:100]}...")
                    is_related_work = True
                    break
            
            if is_related_work:
                continue
            
            is_reference = False
            
            if re.match(r'^\d+\.\s*', line) and not re.match(r'^\[\d+\]', line):
                is_reference = True
                logger.info(f" Found numbered reference: {line[:100]}...")
            
            elif re.search(r'[A-Z][a-z]+,?\s+[A-Z]\.\s*\(\d{4}\)', line):
                is_reference = True
                logger.info(f" Found author-year reference: {line[:100]}...")
            
            elif re.search(r'doi:\s*10\.\d+/', line, re.IGNORECASE):
                is_reference = True
                logger.info(f"Found DOI reference: {line[:100]}...")
            
            elif re.search(r'https?://', line):
                is_reference = True
                logger.info(f"Found URL reference: {line[:100]}...")
            
            elif re.search(r'Journal|Conference|Proceedings|arXiv|IEEE|ACM|Springer', line, re.IGNORECASE):
                is_reference = True
                logger.info(f"Found publication reference: {line[:100]}...")
            
            elif re.search(r'vol\.?\s*\d+|pp\.?\s*\d+|\d+\(\d+\)', line, re.IGNORECASE):
                is_reference = True
                logger.info(f"Found volume/page reference: {line[:100]}...")
            
            elif re.search(r'published|proceedings|conference|journal|article', line, re.IGNORECASE):
                if not any(phrase in line_lower for phrase in ['this paper', 'our study', 'we present', 'proposes']):
                    is_reference = True
                    logger.info(f" Found publication info reference: {line[:100]}...")
            
            elif re.search(r'\b(19|20)\d{2}\b', line) and len(line) > 30:  # Lowered threshold
                if re.search(r'[A-Z][a-z]+|doi|journal|conference|proceedings|http', line, re.IGNORECASE):
                    is_reference = True
                    logger.info(f" Found year-based reference: {line[:100]}...")
            
            if is_reference:
                potential_refs.append(line)
        
        logger.info(f" Aggressive extraction found {len(potential_refs)} potential references")
        
        for i, ref_text in enumerate(potential_refs):
            references.append({
                "raw": ref_text.strip()
            })
        
        return references
    
    def _extract_references_loose(self, text: str) -> List[Dict[str, Any]]:
        """Very loose reference extraction for cases where section headers are removed"""
        references = []
        
        logger.info(" Using very loose reference extraction")
        
        lines = text.split('\n')
        
        potential_refs = []
        
        for line in lines:
            line = line.strip()
            if not line or len(line) < 15:  # Very low threshold
                continue
            
            line_lower = line.lower()
            
            if any(indicator in line_lower for indicator in [
                'related works', 'related work', 'literature review', 'background',
                'introduction', 'methodology', 'conclusion', 'discussion'
            ]):
                continue
            
            is_reference = False
            
            if re.match(r'^\d+\.\s*', line):
                is_reference = True
                logger.info(f"Found loose numbered reference: {line[:100]}...")
            
            elif re.search(r'https?://', line):
                is_reference = True
                logger.info(f"Found loose URL reference: {line[:100]}...")
            
            elif re.search(r'\b(19|20)\d{2}\b', line) and len(line) > 25:
                is_reference = True
                logger.info(f"Found loose year-based reference: {line[:100]}...")
            
            elif re.search(r'[A-Z][a-z]+,\s*[A-Z]', line):
                is_reference = True
                logger.info(f"Found loose author reference: {line[:100]}...")
            
            if is_reference:
                potential_refs.append(line)
        
        logger.info(f"Loose extraction found {len(potential_refs)} potential references")
        
        for i, ref_text in enumerate(potential_refs):
            references.append({
                "raw": ref_text.strip()
            })
        
        return references
    
    def _extract_paper_metadata(self, doc_path: str) -> Dict[str, Any]:
        """Extract basic paper metadata from Word document"""
        metadata = {
            "title": "",
            "authors": [],
            "abstract": "",
            "keywords": [],
            "pages": 0
        }
        
        try:
            doc = Document(doc_path)
            
            metadata["pages"] = len(doc.paragraphs) // 50  # Rough estimate
            
            for i, paragraph in enumerate(doc.paragraphs[:10]):  # Check first 10 paragraphs
                text = paragraph.text.strip()
                
                if text and not metadata["title"]:
                    if len(text) > 10 and len(text) < 200:
                        metadata["title"] = text
                
                if text and not metadata["abstract"]:
                    if 'abstract' in text.lower() and len(text) > 50:
                        metadata["abstract"] = text[:500]  # First 500 chars
                        
        except Exception as e:
            logger.warning(f"Metadata extraction failed: {e}")
        
        return metadata
    
    def detect_paper_type(self, doc_path: str) -> str:
        """Detect paper type based on content"""
        try:
            doc = Document(doc_path)
            
            text = ""
            for paragraph in doc.paragraphs[:10]:
                text += paragraph.text + " "
            
            text_lower = text.lower()
                
            if any(conf in text_lower for conf in ['acl', 'emnlp', 'naacl', 'conll']):
                return "ACL"
            elif any(conf in text_lower for conf in ['aaai', 'ijcai', 'icml', 'iclr', 'neurips']):
                return "AAAI"
            elif any(conf in text_lower for conf in ['ieee', 'acm', 'springer']):
                return "IEEE"
            else:
                return "Generic"
                
        except Exception as e:
            logger.warning(f"Paper type detection failed: {e}")
            return "Generic"
    
    def get_supported_paper_types(self) -> List[str]:
        """Get list of supported paper types"""
        return ["AAAI", "ACL", "ICML", "ICLR", "NeurIPS", "IEEE", "ACM", "Springer", "Generic"]
